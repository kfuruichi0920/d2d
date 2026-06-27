"""Word (.docx) 抽出コマンド

抽出対象: 章・節（見出し）、段落、表、図（blob参照）、箇条書き、脚注
出力形式: {"items": [...], "structure_json": {...}}
"""

from __future__ import annotations
import os
import uuid
from typing import Any, Callable


def run(params: dict, progress: Callable[[str], None]) -> dict:
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    blob_path: str = params["blob_path"]
    progress(f"Word 文書を読み込み中: {os.path.basename(blob_path)}")

    doc = docx.Document(blob_path)
    items: list[dict[str, Any]] = []
    structure: list[dict] = []
    heading_stack: list[dict] = []

    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except ValueError:
                level = 1

            item = {
                "item_uid": str(uuid.uuid4()),
                "item_type": "label",
                "level": level,
                "text": text,
                "style": style_name,
                "para_index": i,
            }
            items.append(item)
            node = {"uid": item["item_uid"], "title": text, "level": level, "children": []}

            # heading_stack の深さを調整
            while heading_stack and heading_stack[-1]["level"] >= level:
                heading_stack.pop()
            if heading_stack:
                heading_stack[-1]["children"].append(node)
            else:
                structure.append(node)
            heading_stack.append(node)

        else:
            items.append({
                "item_uid": str(uuid.uuid4()),
                "item_type": "text",
                "text": text,
                "style": style_name,
                "para_index": i,
            })

    # 表
    for t_idx, table in enumerate(doc.tables):
        progress(f"  表 {t_idx + 1} / {len(doc.tables)} を抽出中")
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        items.append({
            "item_uid": str(uuid.uuid4()),
            "item_type": "table",
            "table_index": t_idx,
            "rows": rows,
            "row_count": len(rows),
            "column_count": len(rows[0]) if rows else 0,
        })

    progress(f"抽出完了: {len(items)} アイテム")
    return {"items": items, "structure_json": structure}
